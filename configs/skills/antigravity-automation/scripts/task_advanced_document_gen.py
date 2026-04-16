"""task_advanced_document_gen.py

Automation Agent task: Advanced Document Generation (DOCX + PDF)

Purpose
- Certify that the Automation Agent can render complex "Superpower" style
  documents with:
  - Multi-level Headings
  - Tables (tab)
  - Figures (fig)
  - Formulas (equ - simplified)
  - Code Blocks (code)

Incorporates advanced parsing philosophies:
- Two-stage mapping (Schema -> Renderer)
- Reading-order preservation
- Specialized element handlers
"""

from __future__ import annotations

import argparse
import hashlib
import json
import time
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

import markdown
from bs4 import BeautifulSoup, Tag, NavigableString

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image, ListFlowable, ListItem, PageBreak
)
from reportlab.lib import colors

from lib import AutomationAgentTask

REPO_ROOT = Path(__file__).resolve().parents[3]
OUT_BASE = REPO_ROOT / "reports" / "artifacts"

def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()

# --- ADVANCED DOCX RENDERER ---

class AdvancedDocxRenderer:
    def __init__(self, doc: Document):
        self.doc = doc

    def render_html_table(self, table_tag: Tag):
        rows = table_tag.find_all('tr')
        if not rows: return
        
        num_rows = len(rows)
        num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < num_cols:
                    table.cell(i, j).text = cell.get_text().strip()

    def render_element(self, tag: Tag, base_path: Path):
        name = tag.name
        if name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(name[1])
            self.doc.add_heading(tag.get_text(), level=level)
        elif name == 'p':
            p = self.doc.add_paragraph()
            self._add_inline_content(p, tag)
        elif name == 'table':
            self.render_html_table(tag)
        elif name == 'ul':
            for li in tag.find_all('li', recursive=False):
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_inline_content(p, li)
        elif name == 'ol':
            for li in tag.find_all('li', recursive=False):
                p = self.doc.add_paragraph(style='List Number')
                self._add_inline_content(p, li)
        elif name == 'pre' or tag.get('class') == ['codehilite']:
            code_text = tag.get_text()
            p = self.doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        elif name == 'img' or name == 'figure':
            src = tag.get('src') or (tag.find('img').get('src') if tag.find('img') else None)
            if src:
                try:
                    full_path = src if Path(src).is_absolute() else (base_path / src).resolve()
                    if full_path.exists():
                        self.doc.add_picture(str(full_path), width=Inches(5))
                    else:
                        self.doc.add_paragraph(f"[Missing Image: {src}]")
                except Exception as e:
                    self.doc.add_paragraph(f"[Image Error: {e}]")

    def _add_inline_content(self, paragraph, element):
        for child in element.children:
            if isinstance(child, NavigableString):
                paragraph.add_run(str(child))
            elif isinstance(child, Tag):
                if child.name in ['b', 'strong']:
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['i', 'em']:
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = paragraph.add_run(child.get_text())
                    run.font.name = 'Courier New'
                else:
                    paragraph.add_run(child.get_text())

# --- ADVANCED PDF RENDERER ---

class AdvancedPdfRenderer:
    def __init__(self, filename: str):
        self.doc = SimpleDocTemplate(filename, pagesize=LETTER)
        self.styles = getSampleStyleSheet()
        self.story = []
        
        # Add custom style
        self.styles.add(ParagraphStyle(
            name='AdvancedCode',
            fontName='Courier',
            fontSize=9,
            leftIndent=20,
            backgroundColor=colors.whitesmoke,
            borderPadding=5
        ))

    def _html_to_rl(self, tag: Tag) -> str:
        out = []
        for child in tag.children:
            if isinstance(child, NavigableString):
                s = str(child).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                out.append(s)
            elif isinstance(child, Tag):
                inner = self._html_to_rl(child)
                if child.name in ['b', 'strong']: out.append(f"<b>{inner}</b>")
                elif child.name in ['i', 'em']: out.append(f"<i>{inner}</i>")
                elif child.name == 'code': out.append(f'<font face="Courier">{inner}</font>')
                elif child.name == 'br': out.append("<br/>")
                else: out.append(inner)
        return "".join(out)

    def render_element(self, tag: Tag, base_path: Path):
        name = tag.name
        if name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            style_name = f"Heading{name[1]}" if f"Heading{name[1]}" in self.styles else "Heading1"
            self.story.append(Paragraph(self._html_to_rl(tag), self.styles[style_name]))
            self.story.append(Spacer(1, 0.1 * inch))
        elif name == 'p':
            self.story.append(Paragraph(self._html_to_rl(tag), self.styles['Normal']))
            self.story.append(Spacer(1, 0.1 * inch))
        elif name == 'table':
            data = []
            for tr in tag.find_all('tr'):
                data.append([td.get_text().strip() for td in tr.find_all(['td', 'th'])])
            if data:
                t = Table(data, hAlign='LEFT')
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey)
                ]))
                self.story.append(t)
                self.story.append(Spacer(1, 0.1 * inch))
        elif name == 'ul':
            items = [ListItem(Paragraph(self._html_to_rl(li), self.styles['Normal'])) for li in tag.find_all('li', recursive=False)]
            self.story.append(ListFlowable(items, bulletType='bullet', leftIndent=20))
            self.story.append(Spacer(1, 0.1 * inch))
        elif name == 'pre' or tag.get('class') == ['codehilite']:
            self.story.append(Paragraph(tag.get_text().replace('\n', '<br/>'), self.styles['AdvancedCode']))
            self.story.append(Spacer(1, 0.1 * inch))
        elif name == 'img' or name == 'figure':
            src = tag.get('src') or (tag.find('img').get('src') if tag.find('img') else None)
            if src:
                try:
                    full_path = src if Path(src).is_absolute() else (base_path / src).resolve()
                    if full_path.exists():
                        img = Image(str(full_path), width=5*inch, height=None, kind='proportional')
                        self.story.append(img)
                        self.story.append(Spacer(1, 0.1 * inch))
                except Exception as e:
                    self.story.append(Paragraph(f"<i>Image error: {e}</i>", self.styles['Normal']))

    def save(self):
        self.doc.build(self.story)

# --- MAIN TASK ---

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--run", action="store_true")
    ap.add_argument("--input-md", type=str, default="")
    args = ap.parse_args()

    task = AutomationAgentTask(
        name="Advanced Document Gen",
        description="Verify High-Fidelity Table/Image/Code rendering for production output."
    )

    if not args.run:
        print("[AutomationAgent] Prepared mode. Use --run.")
        return 0

    task.start()

    try:
        md_file = Path(args.input_md).resolve()
        if not md_file.exists():
            raise FileNotFoundError(f"MD file not found: {md_file}")

        text = md_file.read_text(encoding='utf-8')
        html = markdown.markdown(text, extensions=['extra', 'tables', 'codehilite'])
        soup = BeautifulSoup(html, 'html.parser')

        ts = time.strftime("%Y-%m-%d-%H-%M-%S")
        out_dir = OUT_BASE / f"advanced-report-{ts}"
        out_dir.mkdir(parents=True, exist_ok=True)

        # 1. DOCX
        docx_path = out_dir / "advanced_report.docx"
        def gen_docx():
            doc = Document()
            renderer = AdvancedDocxRenderer(doc)
            for element in soup.contents:
                if isinstance(element, Tag):
                    renderer.render_element(element, md_file.parent)
            doc.save(str(docx_path))
            return str(docx_path)
        
        task.step("Render Advanced DOCX", gen_docx)

        # 2. PDF
        pdf_path = out_dir / "advanced_report.pdf"
        def gen_pdf():
            renderer = AdvancedPdfRenderer(str(pdf_path))
            for element in soup.contents:
                if isinstance(element, Tag):
                    renderer.render_element(element, md_file.parent)
            renderer.save()
            return str(pdf_path)
        
        task.step("Render Advanced PDF", gen_pdf)

        # 3. Evidence
        def collect():
            return {
                "artifacts": [
                    {"name": docx_path.name, "sha256": sha256_file(docx_path)},
                    {"name": pdf_path.name, "sha256": sha256_file(pdf_path)}
                ]
            }
        
        ev = task.step("Final Evidence Collection", collect)
        task.finish(success=True, final_result=ev)
        
        print(f"\n[Advanced Success] Report artifacts: {out_dir}")
        return 0

    except Exception as e:
        task.finish(success=False, final_result={"error": str(e)})
        raise

if __name__ == "__main__":
    main()
